"""Regression tests for real-world resume formats and malformed documents.

Each test here corresponds to a failure observed while parsing genuine
PDF/DOCX layouts, so the suite doubles as a bug-regression record.
"""

from __future__ import annotations

import io

import pytest

from resume_analyzer.exceptions import (
    CorruptDocumentError,
    EmptyDocumentError,
    UnsupportedFileTypeError,
)
from resume_analyzer.parsing.document import extract_document
from resume_analyzer.parsing.resume_parser import (
    _titlecase_name,
    education_level,
    extract_phone,
    parse_resume,
    split_sections,
)


class FakeUpload(io.BytesIO):
    """Stand-in for a Streamlit ``UploadedFile``."""

    def __init__(self, payload: bytes, name: str):
        super().__init__(payload)
        self.name = name
        self.size = len(payload)


@pytest.fixture(scope="module")
def text_pdf_bytes() -> bytes:
    """A single-page, text-based PDF resume."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    y = 750
    for line in [
        "Muhammad Awais Khan",
        "Machine Learning Engineer",
        "awais.khan@gmail.com | +92 300 1234567 | Lahore, Pakistan",
        "linkedin.com/in/awaiskhan | github.com/awais18",
        "",
        "EXPERIENCE",
        "Machine Learning Engineer | Systems Ltd    Mar 2021 - Present",
        "- Deployed BERT-based classifier improving accuracy by 18%.",
        "",
        "EDUCATION",
        "BS in Computer Science, FAST NUCES (2019)",
        "",
        "SKILLS",
        "Python, TensorFlow, PyTorch, NLP, SQL, Docker, AWS",
        "",
        "LANGUAGES",
        "English, Urdu, Punjabi",
    ]:
        pdf.drawString(60, y, line)
        y -= 15
    pdf.save()
    return buffer.getvalue()


@pytest.fixture(scope="module")
def scanned_pdf_bytes() -> bytes:
    """A PDF with no text layer, simulating a scanned resume."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


@pytest.fixture(scope="module")
def table_docx_bytes() -> bytes:
    """A DOCX using a table-based layout, as produced by many CV templates."""
    from docx import Document

    document = Document()
    document.add_paragraph("Sarah Connor-Smith")
    document.add_paragraph("sarah.oconnor@company.co.uk | +44 20 7946 0958")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "SKILLS"
    table.cell(0, 1).text = "Java, Spring Boot, Kubernetes, PostgreSQL"
    table.cell(1, 0).text = "EDUCATION"
    table.cell(1, 1).text = "MSc Software Engineering, Oxford 2017"
    document.add_paragraph("EXPERIENCE")
    document.add_paragraph("Senior Java Developer at Barclays  2018 - Present")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --------------------------------------------------------------------------
# Real PDF parsing
# --------------------------------------------------------------------------


def test_parses_text_pdf(text_pdf_bytes: bytes) -> None:
    document = extract_document(FakeUpload(text_pdf_bytes, "awais.pdf"))
    profile = parse_resume(document.text)

    assert document.page_count == 1
    assert profile.contact.full_name == "Muhammad Awais Khan"
    assert profile.contact.email == "awais.khan@gmail.com"
    assert profile.contact.phone is not None
    assert "linkedin.com/in/awaiskhan" in (profile.contact.linkedin or "")
    assert "github.com/awais18" in (profile.contact.github or "")


def test_pdf_education_and_skills(text_pdf_bytes: bytes) -> None:
    profile = parse_resume(extract_document(FakeUpload(text_pdf_bytes, "a.pdf")).text)

    assert profile.education
    assert profile.education[0].degree == "BS"
    assert "Python" in profile.skill_names
    assert "Urdu" in profile.languages


def test_scanned_pdf_raises_friendly_error(scanned_pdf_bytes: bytes) -> None:
    with pytest.raises(EmptyDocumentError) as info:
        extract_document(FakeUpload(scanned_pdf_bytes, "scan.pdf"))
    assert "scanned" in info.value.user_message.lower()


# --------------------------------------------------------------------------
# DOCX with table layout
# --------------------------------------------------------------------------


def test_parses_table_based_docx(table_docx_bytes: bytes) -> None:
    document = extract_document(FakeUpload(table_docx_bytes, "sarah.docx"))
    profile = parse_resume(document.text)

    assert profile.contact.email == "sarah.oconnor@company.co.uk"
    assert {"Java", "Kubernetes", "PostgreSQL"} <= set(profile.skill_names)


def test_inline_table_heading_is_routed(table_docx_bytes: bytes) -> None:
    """'SKILLS | Java, ...' must populate the skills section, not the header."""
    text = extract_document(FakeUpload(table_docx_bytes, "s.docx")).text
    sections = split_sections(text)
    assert "skills" in sections
    assert "Java" in sections["skills"]


def test_section_heading_is_not_an_experience_entry(table_docx_bytes: bytes) -> None:
    profile = parse_resume(extract_document(FakeUpload(table_docx_bytes, "s.docx")).text)
    titles = [entry.title.upper() for entry in profile.experience]
    assert "EDUCATION" not in titles


# --------------------------------------------------------------------------
# Phone formats
# --------------------------------------------------------------------------


@pytest.mark.parametrize(
    "raw",
    [
        "+1 (415) 555-0198",
        "+44 20 7946 0958",
        "+92 300 1234567",
        "0300-1234567",
        "415.555.0198",
    ],
)
def test_international_phone_formats(raw: str) -> None:
    assert extract_phone(f"Contact me at {raw} anytime") is not None


@pytest.mark.parametrize("raw", ["2019 - 2022", "Section 4.2.1"])
def test_non_phone_numbers_are_rejected(raw: str) -> None:
    assert extract_phone(raw) is None


# --------------------------------------------------------------------------
# Degree disambiguation
# --------------------------------------------------------------------------


def test_short_degree_abbreviations_are_detected() -> None:
    assert education_level("BS in Computer Science")[0] == "BS"
    assert education_level("MS in Data Science")[0] == "MS"


def test_product_names_are_not_degrees() -> None:
    assert education_level("Proficient in MS Excel and MS Office")[1] == 0


def test_degree_not_matched_inside_words() -> None:
    assert education_level("Managed systems and forms for jobs")[1] == 0


# --------------------------------------------------------------------------
# Name casing
# --------------------------------------------------------------------------


@pytest.mark.parametrize(
    ("raw", "expected"),
    [
        ("JOHN O'CONNOR", "John O'Connor"),
        ("maria ramirez-alvarez", "Maria Ramirez-Alvarez"),
        ("McLeod Stone", "McLeod Stone"),
    ],
)
def test_name_titlecasing(raw: str, expected: str) -> None:
    assert _titlecase_name(raw) == expected


def test_unicode_resume_is_parsed() -> None:
    profile = parse_resume(
        "José Ramírez\njose@example.com\nSKILLS\nPython • Django • Docker"
    )
    assert profile.contact.email == "jose@example.com"
    assert {"Python", "Django", "Docker"} <= set(profile.skill_names)


# --------------------------------------------------------------------------
# Malformed documents
# --------------------------------------------------------------------------


@pytest.mark.parametrize(
    ("payload", "name", "expected"),
    [
        (b"%PDF-1.4 broken garbage", "corrupt.pdf", CorruptDocumentError),
        (b"", "empty.pdf", EmptyDocumentError),
        (b"PK\x03\x04garbage", "fake.docx", CorruptDocumentError),
        (b"plain text", "notes.txt", UnsupportedFileTypeError),
    ],
)
def test_malformed_documents_raise_typed_errors(payload, name, expected) -> None:
    with pytest.raises(expected):
        extract_document(FakeUpload(payload, name))


def test_error_messages_never_leak_internals() -> None:
    for payload, name in [(b"garbage", "x.pdf"), (b"", "y.docx")]:
        try:
            extract_document(FakeUpload(payload, name))
        except Exception as exc:
            message = getattr(exc, "user_message", "")
            assert message
            assert "Traceback" not in message
            assert "Error:" not in message

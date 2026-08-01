"""Document text extraction for PDF and DOCX resumes.

The extractor is defensive by design: uploads are validated before they are
opened, per-page failures never abort the whole document, and unreadable files
raise a typed exception carrying a user-safe message.
"""

from __future__ import annotations

import io
from dataclasses import dataclass
from pathlib import Path
from typing import Any, BinaryIO, Protocol

from resume_analyzer.config.constants import (
    MAX_UPLOAD_BYTES,
    MIN_MEANINGFUL_TEXT_LENGTH,
    SUPPORTED_EXTENSIONS,
)
from resume_analyzer.config.logging_config import get_logger
from resume_analyzer.exceptions import (
    CorruptDocumentError,
    EmptyDocumentError,
    FileTooLargeError,
    UnsupportedFileTypeError,
)
from resume_analyzer.utils_text import normalize_whitespace

logger = get_logger(__name__)


class UploadedDocument(Protocol):
    """Structural type matching Streamlit's ``UploadedFile``."""

    name: str

    def read(self, size: int = -1) -> bytes: ...

    def seek(self, offset: int, whence: int = 0) -> int: ...


@dataclass(frozen=True, slots=True)
class ExtractedDocument:
    """Result of reading a resume file.

    Attributes:
        text: Normalised plain text.
        filename: Original upload name.
        page_count: Number of pages (PDF) or paragraphs+tables (DOCX).
        warnings: Non-fatal issues worth surfacing to the user.
    """

    text: str
    filename: str
    page_count: int = 0
    warnings: tuple[str, ...] = ()

    @property
    def is_empty(self) -> bool:
        """``True`` when the document holds no usable text."""
        return len(self.text.strip()) < MIN_MEANINGFUL_TEXT_LENGTH


def _file_extension(filename: str) -> str:
    """Return the lowercase suffix of ``filename``."""
    return Path(filename or "").suffix.lower()


def _read_bytes(file: Any) -> bytes:
    """Read an upload defensively, rewinding first so re-reads are safe."""
    try:
        if hasattr(file, "seek"):
            file.seek(0)
        data = file.getvalue() if hasattr(file, "getvalue") else file.read()
    except Exception as exc:
        logger.warning("Unable to read uploaded file: %s", exc)
        raise CorruptDocumentError(f"read failed: {exc}") from exc

    if isinstance(data, str):
        data = data.encode("utf-8", errors="ignore")
    if not data:
        raise EmptyDocumentError("uploaded file is empty")
    return data


def validate_upload(file: Any) -> None:
    """Validate an upload's type and size before any parsing work.

    Args:
        file: A Streamlit ``UploadedFile`` or any object with ``name``/``read``.

    Raises:
        UnsupportedFileTypeError: The extension is not PDF or DOCX.
        FileTooLargeError: The payload exceeds the configured limit.
    """
    if file is None:
        raise UnsupportedFileTypeError("no file supplied")

    extension = _file_extension(getattr(file, "name", ""))
    if extension not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(f"unsupported extension: {extension!r}")

    size = getattr(file, "size", None)
    if isinstance(size, int) and size > MAX_UPLOAD_BYTES:
        raise FileTooLargeError(f"{size} bytes exceeds limit")


def extract_pdf_text(source: bytes | BinaryIO) -> tuple[str, int, list[str]]:
    """Extract text from a PDF.

    Falls back to word-level extraction when a page's layout engine returns
    nothing, and tolerates individual page failures.

    Args:
        source: PDF bytes or a binary stream.

    Returns:
        Tuple of ``(text, page_count, warnings)``.

    Raises:
        CorruptDocumentError: The PDF cannot be opened at all.
    """
    import pdfplumber

    stream = io.BytesIO(source) if isinstance(source, bytes) else source
    chunks: list[str] = []
    warnings: list[str] = []

    try:
        with pdfplumber.open(stream) as pdf:
            page_count = len(pdf.pages)
            for index, page in enumerate(pdf.pages, start=1):
                try:
                    page_text = page.extract_text() or ""
                    if not page_text.strip():
                        words = page.extract_words() or []
                        page_text = " ".join(word.get("text", "") for word in words)
                    if page_text.strip():
                        chunks.append(page_text)
                    else:
                        warnings.append(f"Page {index} contains no selectable text.")
                except Exception as exc:  # pragma: no cover - per-page guard
                    logger.warning("PDF page %s failed: %s", index, exc)
                    warnings.append(f"Page {index} could not be read.")
    except Exception as exc:
        logger.error("Unable to open PDF: %s", exc)
        raise CorruptDocumentError(f"pdf open failed: {exc}") from exc

    return "\n".join(chunks), page_count, warnings


def extract_docx_text(source: bytes | BinaryIO) -> tuple[str, int, list[str]]:
    """Extract text from a DOCX file, including tables, headers and footers.

    Args:
        source: DOCX bytes or a binary stream.

    Returns:
        Tuple of ``(text, block_count, warnings)``.

    Raises:
        CorruptDocumentError: The document cannot be opened.
    """
    from docx import Document

    stream = io.BytesIO(source) if isinstance(source, bytes) else source
    chunks: list[str] = []
    warnings: list[str] = []

    try:
        document = Document(stream)
    except Exception as exc:
        logger.error("Unable to open DOCX: %s", exc)
        raise CorruptDocumentError(f"docx open failed: {exc}") from exc

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            chunks.append(paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))

    try:  # Headers and footers often hold contact details.
        for section in document.sections:
            for container in (section.header, section.footer):
                for paragraph in container.paragraphs:
                    if paragraph.text.strip():
                        chunks.append(paragraph.text.strip())
    except Exception as exc:  # pragma: no cover - optional content
        logger.debug("Header/footer extraction skipped: %s", exc)
        warnings.append("Headers and footers could not be read.")

    return "\n".join(chunks), len(chunks), warnings


def extract_document(file: Any) -> ExtractedDocument:
    """Read an uploaded resume into normalised text.

    Args:
        file: Streamlit ``UploadedFile`` or file-like object with a ``name``.

    Returns:
        A populated :class:`ExtractedDocument`.

    Raises:
        UnsupportedFileTypeError: Extension is not supported.
        FileTooLargeError: Upload exceeds the size limit.
        CorruptDocumentError: The file cannot be decoded.
        EmptyDocumentError: No meaningful text could be extracted.
    """
    validate_upload(file)

    filename = getattr(file, "name", "resume")
    extension = _file_extension(filename)
    payload = _read_bytes(file)

    if len(payload) > MAX_UPLOAD_BYTES:
        raise FileTooLargeError(f"{len(payload)} bytes exceeds limit")

    if extension == ".pdf":
        raw_text, page_count, warnings = extract_pdf_text(payload)
    else:
        raw_text, page_count, warnings = extract_docx_text(payload)

    document = ExtractedDocument(
        text=normalize_whitespace(raw_text),
        filename=filename,
        page_count=page_count,
        warnings=tuple(warnings),
    )

    if document.is_empty:
        logger.warning("No usable text extracted from %s", filename)
        raise EmptyDocumentError(f"no text extracted from {filename}")

    logger.info(
        "Extracted %s characters from %s (%s pages).",
        len(document.text),
        filename,
        page_count,
    )
    return document
